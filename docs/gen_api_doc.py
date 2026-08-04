# -*- coding: utf-8 -*-
"""
扫描 backend/**/controller/*.java，提取接口信息，生成 Word 文档。
输出：docs/关键接口解析.docx
"""
import os
import re
import glob
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.abspath(__file__))
BACKEND = os.path.join(ROOT, "..", "backend")
OUTPUT = os.path.join(ROOT, "关键接口解析.docx")

# 接口映射注解 -> HTTP 方法
METHOD_ANN = {
    "@GetMapping": "GET",
    "@PostMapping": "POST",
    "@PutMapping": "PUT",
    "@DeleteMapping": "DELETE",
    "@PatchMapping": "PATCH",
}

# 控制器中文化映射
MODULE_CN = {
    "AuthController": "认证与登录",
    "SysUserController": "用户管理",
    "SysRoleController": "角色管理",
    "SysMenuController": "菜单管理",
    "SysDeptController": "部门管理",
    "SysLogController": "操作日志",
    "SysMessageController": "系统消息",
    "IntegrationController": "第三方集成",
    "FileController": "文件上传",
    "ClueController": "线索管理",
    "ChannelController": "渠道管理",
    "KnowledgeController": "知识库",
    "ExternalClueController": "外部线索接入",
    "CustomerController": "客户管理",
    "ContactController": "联系人管理",
    "TagController": "标签管理",
    "FollowRecordController": "跟进记录",
    "OpportunityController": "商机管理",
    "ContractController": "合同管理",
    "PaymentController": "回款管理",
    "SignInController": "外勤签到",
    "ServiceOrderController": "工单管理",
    "ServiceRecordController": "售后记录",
    "ReportController": "数据分析与报表",
}


def parse_controller(filepath):
    """解析单个 Controller 文件，返回 (类名, 模块中文, RequestMapping 前缀, [接口列表])"""
    with open(filepath, encoding="utf-8") as f:
        content = f.read()

    # 类名
    m = re.search(r"class\s+(\w+Controller)", content)
    if not m:
        return None
    class_name = m.group(1)
    module_cn = MODULE_CN.get(class_name, class_name)

    # 类级 @RequestMapping
    m = re.search(r'@RequestMapping\(\s*"([^"]*)"\s*\)', content)
    class_prefix = m.group(1) if m else ""

    # 逐行状态机解析，避免正则跨越 @Operation 导致错配
    apis = []
    pending_summary = None  # 最近一次遇到的 @Operation summary
    pending_http_method = None
    pending_http_path = None

    # 用正则逐块扫描：先找所有注解和方法签名的位置
    token_pattern = re.compile(
        r'@Operation\(summary\s*=\s*"([^"]*)"\)'   # 1: summary
        r'|(@GetMapping|@PostMapping|@PutMapping|@DeleteMapping|@PatchMapping)'  # 2: HTTP 注解
        r'(?:\(\s*(?:value\s*=\s*)?"([^"]*)"\s*\))?'  # 3: 路径（可选）
        r'|public\s+[\w<>,\s\[\]\?]+\s+(\w+)\s*\('  # 4: 方法名
        r'([^)]*)\)',  # 5: 参数列表
        re.DOTALL
    )

    for match in token_pattern.finditer(content):
        if match.group(1) is not None:
            # 遇到 @Operation
            pending_summary = match.group(1)
        elif match.group(2) is not None:
            # 遇到 HTTP 注解
            pending_http_method = METHOD_ANN.get(match.group(2), "?")
            pending_http_path = match.group(3) or ""
        elif match.group(4) is not None:
            # 遇到 public 方法签名
            method_name = match.group(4)
            params_raw = match.group(5).strip() if match.group(5) else ""
            # 只有在 method_name 之前有 summary 和 http_method 才算接口
            if pending_http_method is not None:
                summary = pending_summary or ""
                path = pending_http_path or ""
                full_path = (class_prefix + path).replace("//", "/")
                if not full_path.startswith("/"):
                    full_path = "/" + full_path

                # 解析参数
                params = []
                if params_raw:
                    for pm in re.finditer(r'@(PathVariable|RequestParam)(?:\(\s*(?:value\s*=\s*)?"[^"]*"\s*\))?\s+(\w+)\s+(\w+)', params_raw):
                        params.append(f"{pm.group(3)} ({pm.group(2)})")
                    for pm in re.finditer(r'@RequestBody\s+(\w+)\s+(\w+)', params_raw):
                        params.append(f"{pm.group(2)} ({pm.group(1)})")
                    if not params:
                        for pm in re.finditer(r'\b(\w+DTO|\w+PageDTO|\w+VO)\s+(\w+)', params_raw):
                            params.append(f"{pm.group(2)} ({pm.group(1)})")

                apis.append({
                    "summary": summary,
                    "method": pending_http_method,
                    "path": full_path,
                    "java_method": method_name,
                    "params": params,
                })
            # 重置
            pending_summary = None
            pending_http_method = None
            pending_http_path = None

    return {
        "class_name": class_name,
        "module_cn": module_cn,
        "prefix": class_prefix,
        "file_path": filepath,
        "apis": apis,
    }


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_api_table(doc, controller):
    """为控制器添加接口表格"""
    if not controller["apis"]:
        doc.add_paragraph("（未扫描到接口）")
        return

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr = table.rows[0].cells
    headers = ["HTTP", "路径", "说明", "Java 方法", "参数"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(hdr[i], "4472C4")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 数据行
    for api in controller["apis"]:
        row = table.add_row().cells
        row[0].text = api["method"]
        row[1].text = api["path"]
        row[2].text = api["summary"]
        row[3].text = api["java_method"]
        row[4].text = "\n".join(api["params"]) if api["params"] else "-"
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # 设置列宽
    widths = [Inches(0.6), Inches(2.2), Inches(1.5), Inches(1.2), Inches(1.8)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


def main():
    # 找到所有 Controller
    files = glob.glob(os.path.join(BACKEND, "**", "controller", "*.java"), recursive=True)
    files.sort()

    controllers = []
    for f in files:
        if os.path.basename(f) == "CrmApplication.java":
            continue
        c = parse_controller(f)
        if c:
            controllers.append(c)

    # 按模块中文分组排序
    controllers.sort(key=lambda x: x["module_cn"])

    # 创建文档
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # 标题
    title = doc.add_heading("CRM 系统 - 关键接口解析与位置", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 说明
    p = doc.add_paragraph()
    run = p.add_run("本文档由代码扫描自动生成，涵盖后端全部 Controller 的接口清单、HTTP 方法、请求路径、Java 方法名及参数。")
    run.font.size = Pt(10)
    run.italic = True

    doc.add_paragraph()

    # 目录统计
    doc.add_heading("一、接口总览", level=1)
    total = sum(len(c["apis"]) for c in controllers)
    p = doc.add_paragraph()
    p.add_run(f"共扫描到 ").font.size = Pt(11)
    r = p.add_run(f"{len(controllers)} 个 Controller")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.add_run(f"，合计 ").font.size = Pt(11)
    r = p.add_run(f"{total} 个接口")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.add_run("。").font.size = Pt(11)

    # 总览表
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["模块", "Controller", "路径前缀", "接口数"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "4472C4")

    for c in controllers:
        row = table.add_row().cells
        row[0].text = c["module_cn"]
        row[1].text = c["class_name"]
        row[2].text = c["prefix"] or "-"
        row[3].text = str(len(c["apis"]))
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_page_break()

    # 详细解析
    doc.add_heading("二、接口详细解析", level=1)

    for idx, c in enumerate(controllers, 1):
        # 模块标题
        h = doc.add_heading(f"{idx}. {c['module_cn']}（{c['class_name']}）", level=2)

        # 文件位置
        rel_path = os.path.relpath(c["file_path"], ROOT)
        p = doc.add_paragraph()
        r = p.add_run("文件位置：")
        r.bold = True
        r.font.size = Pt(9)
        r2 = p.add_run(rel_path.replace("\\", "/"))
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
        r2.font.name = "Consolas"

        # 路径前缀
        if c["prefix"]:
            p = doc.add_paragraph()
            r = p.add_run("路径前缀：")
            r.bold = True
            r.font.size = Pt(9)
            r2 = p.add_run(c["prefix"])
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            r2.font.name = "Consolas"

        # 接口表
        add_api_table(doc, c)
        doc.add_paragraph()

    # 附录
    doc.add_page_break()
    doc.add_heading("三、附录：HTTP 状态码说明", level=1)
    codes = [
        ("200", "请求成功"),
        ("400", "参数错误 / 业务校验失败"),
        ("401", "未认证（token 缺失或失效）"),
        ("403", "权限不足"),
        ("404", "资源不存在"),
        ("500", "服务器内部错误"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["状态码", "说明"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "4472C4")
    for code, desc in codes:
        row = table.add_row().cells
        row[0].text = code
        row[1].text = desc

    # 统一返回体说明
    doc.add_heading("四、附录：统一返回体 R 格式", level=1)
    p = doc.add_paragraph()
    p.add_run("所有接口统一返回 ").font.size = Pt(10)
    r = p.add_run("com.crm.common.api.R")
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    p.add_run("，结构如下：").font.size = Pt(10)

    code_block = doc.add_paragraph()
    code_block.paragraph_format.left_indent = Inches(0.3)
    r = code_block.add_run(
        '{\n'
        '  "code": 200,           // 状态码：200成功 / 400失败 / 401未认证 / 403无权限\n'
        '  "msg": "操作成功",      // 提示消息\n'
        '  "data": {}             // 业务数据\n'
        '}'
    )
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 保存
    doc.save(OUTPUT)
    print(f"已生成: {OUTPUT}")
    print(f"控制器数: {len(controllers)}, 接口总数: {total}")


if __name__ == "__main__":
    main()
